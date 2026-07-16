from pathlib import Path

from tts_app.readers import load_text


def test_read_docx_bulgarian(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Здравей, това е български текст.")
    doc.add_paragraph("Hello English too.")
    doc.save(path)
    text = load_text(str(path))
    assert "български" in text
    assert "Hello" in text


def test_read_pdf_bulgarian(tmp_path: Path) -> None:
    import fitz

    path = tmp_path / "sample.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "България / Bulgaria")
    pdf.save(path)
    pdf.close()
    text = load_text(str(path))
    assert "България" in text or "Bulgaria" in text


def test_read_txt_cp1251(tmp_path: Path) -> None:
    path = tmp_path / "bg.txt"
    path.write_bytes("Здравей".encode("cp1251"))
    assert "Здравей" in load_text(str(path))
